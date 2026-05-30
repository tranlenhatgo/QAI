"""
Generate a Word document from all thesis markdown files.
Preserves formatting: headings, bold, italic, code blocks, tables, lists.
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Thesis files in order
THESIS_FILES = [
    "chapter-01-introduction.md",
    "chapter-02-literature-review.md",
    "chapter-03-requirements-analysis.md",
    "chapter-04-system-design-part1.md",
    "chapter-04-system-design-part2.md",
    "chapter-04-system-design-part3.md",
    "chapter-05-implementation-part1.md",
    "chapter-05-implementation-part2.md",
    "chapter-05-implementation-part3.md",
    "chapter-06-testing-evaluation.md",
    "chapter-07-conclusion.md",
    "references.md",
    "appendix-a-architecture-diagrams.md",
    "appendix-b-api-documentation.md",
    "appendix-c-database-schema.md",
]

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))


def set_cell_shading(cell, color):
    """Set background shading of a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_code_paragraph(doc, text):
    """Add a code block paragraph with monospace font and gray background."""
    para = doc.add_paragraph()
    para.style = doc.styles['No Spacing'] if 'No Spacing' in [s.name for s in doc.styles] else None
    para.paragraph_format.left_indent = Cm(1)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    # Add shading
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F5F5F5')
    rPr.append(shd)
    return para


def add_formatted_runs(paragraph, text):
    """Add runs with bold, italic, code formatting."""
    # Pattern matches: **bold**, *italic*, `code`, and plain text
    pattern = r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            # gray background for inline code
            rPr = run._r.get_or_add_rPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), 'E8E8E8')
            rPr.append(shd)
        else:
            paragraph.add_run(part)


def parse_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (rows, end_idx)."""
    rows = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        row_text = lines[i].strip()
        # Split by | and strip
        cells = [c.strip() for c in row_text.split('|')]
        # Remove empty first and last (from leading/trailing |)
        if cells and cells[0] == '':
            cells = cells[1:]
        if cells and cells[-1] == '':
            cells = cells[:-1]
        rows.append(cells)
        i += 1
    return rows, i


def add_table_to_doc(doc, rows):
    """Add a formatted table to the document."""
    if len(rows) < 2:
        return
    
    # Skip separator row (row with ---)
    header = rows[0]
    data_rows = []
    for row in rows[1:]:
        if all(set(cell.strip()) <= set('-| :') for cell in row):
            continue  # skip separator
        data_rows.append(row)
    
    num_cols = len(header)
    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.style = 'Table Grid'
    
    # Header row
    for j, cell_text in enumerate(header):
        if j < num_cols:
            cell = table.rows[0].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.bold = True
            run.font.size = Pt(9)
            set_cell_shading(cell, 'D9E2F3')
    
    # Data rows
    for i, row in enumerate(data_rows):
        for j, cell_text in enumerate(row):
            if j < num_cols:
                cell = table.rows[i + 1].cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                add_formatted_runs(p, cell_text)
                for run in p.runs:
                    run.font.size = Pt(9)
    
    doc.add_paragraph()  # spacing after table


def process_markdown_file(doc, filepath):
    """Process a markdown file and add its content to the Word document."""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # Code block start/end
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block - write collected code
                code_text = '\n'.join(code_lines)
                if code_text.strip():
                    add_code_paragraph(doc, code_text)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Empty line
        if not line.strip():
            i += 1
            continue
        
        # Headings
        if line.startswith('# '):
            heading = doc.add_heading(line[2:].strip(), level=1)
            heading.runs[0].font.size = Pt(18)
            i += 1
            continue
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue
        
        # Table
        if line.strip().startswith('|'):
            rows, end_idx = parse_table(lines, i)
            add_table_to_doc(doc, rows)
            i = end_idx
            continue
        
        # Horizontal rule
        if line.strip() in ('---', '***', '___'):
            doc.add_paragraph('_' * 60)
            i += 1
            continue
        
        # Bullet lists
        bullet_match = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            text = bullet_match.group(2)
            para = doc.add_paragraph()
            para.style = 'List Bullet'
            if indent >= 2:
                para.paragraph_format.left_indent = Cm(1.5 + (indent // 2) * 0.7)
            add_formatted_runs(para, text)
            i += 1
            continue
        
        # Numbered lists
        num_match = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if num_match:
            text = num_match.group(2)
            para = doc.add_paragraph()
            para.style = 'List Number'
            add_formatted_runs(para, text)
            i += 1
            continue
        
        # Regular paragraph
        para = doc.add_paragraph()
        add_formatted_runs(para, line)
        i += 1


def create_thesis_document():
    """Create the complete thesis Word document."""
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)
    
    # Title page
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('GRADUATION THESIS')
    run.bold = True
    run.font.size = Pt(24)
    
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('QAI — AI-Enhanced Quiz and Learning Platform')
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Process each file
    for filename in THESIS_FILES:
        filepath = os.path.join(SCRIPT_DIR, filename)
        if os.path.exists(filepath):
            print(f"Processing: {filename}")
            # Add page break between chapters
            doc.add_page_break()
            process_markdown_file(doc, filepath)
        else:
            print(f"WARNING: File not found: {filename}")
    
    # Save
    output_path = os.path.join(SCRIPT_DIR, "QAI_Graduation_Thesis.docx")
    doc.save(output_path)
    print(f"\n✅ Word document saved: {output_path}")
    return output_path


if __name__ == '__main__':
    create_thesis_document()
